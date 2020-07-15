########################################################################################################################
# Defines a class that handles writing results to a word document.
#
# Written by Rian Koja to publish in a GitHub repository with specified licence.
########################################################################################################################
import os
import io
import getpass
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt


class ReportDocument:
    def __init__(self, title='Post Report', file_name_prefix='Report_v', user_name=None):
        self.document = Document()
        self.document.add_heading(title, level=0)
        if user_name is None:
            user_name = getpass.getuser()
        self.document.add_heading('Author: ' + user_name, level=1)

        self.file_path = os.path.realpath(os.path.join(os.path.dirname(os.path.realpath(__file__)), '..', 'mount'))

        file_version = 0
        while os.path.isfile(os.path.join(self.file_path, file_name_prefix + str(file_version) + ".docx")):
            file_version += 1

        self.file_name = file_name_prefix + str(file_version) + ".docx"

    def add_fig(self, wid=6):
        memfile = io.BytesIO()
        plt.savefig(memfile)
        self.document.add_picture(memfile, width=Inches(wid))
        memfile.close()

    def finish(self):
        self.document.save(os.path.join(self.file_path, self.file_name))
        print("Finished report:", self.file_name)

    def add_heading(self, text, level=2):
        self.document.add_heading(text, level=level)

    def add_paragraph(self, text):
        self.document.add_paragraph(text)


if __name__ == "__main__":
    # Just test the functions:
    testDoc = ReportDocument()
    testDoc.finish()
